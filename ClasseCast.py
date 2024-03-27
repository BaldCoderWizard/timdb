import requests
from docx import Document
from docx.shared import Inches
import os
from docx2pdf import convert

class Elenco:
    def obter_elenco(self, movie_id):
        url = f"https://api.themoviedb.org/3/movie/{movie_id}/credits?language=en-US"
        headers = {"accept": "application/json", "Authorization": "Bearer eyJhbGciOiJIUzI1NiJ9.eyJhdWQiOiIwZGMxMTg3NjU4NjIxNmQ1NWM3ZTQ2Y2EzZjU1ZmJmYiIsInN1YiI6IjY1ZjE3MjFjNmRlYTNhMDE2Mzc4YThjYyIsInNjb3BlcyI6WyJhcGlfcmVhZCJdLCJ2ZXJzaW9uIjoxfQ.mniwNJ0_EVyiz75FjAgZWI2SUFKoTV2A9eci3cpcEhQ"}       
        response = requests.get(url, headers=headers)
        return response.json()

    def selecionar_nome_foto_atores(self, elenco_json):
        url_base_imagem = "https://image.tmdb.org/t/p/w500"
    
        atores_info = []
    
        if 'cast' in elenco_json and elenco_json['cast']:
            for ator in elenco_json['cast'][:5]:
                url_foto = url_base_imagem + ator['profile_path'] if ator.get('profile_path') else None
                atores_info.append({'nome': ator['name'], 'url_foto': url_foto})
    
        return atores_info

    def obter_detalhes_filme(self, movie_id):
        url = f"https://api.themoviedb.org/3/movie/{movie_id}?language=en-US"
        headers = {"accept": "application/json", "Authorization": "Bearer eyJhbGciOiJIUzI1NiJ9.eyJhdWQiOiIwZGMxMTg3NjU4NjIxNmQ1NWM3ZTQ2Y2EzZjU1ZmJmYiIsInN1YiI6IjY1ZjE3MjFjNmRlYTNhMDE2Mzc4YThjYyIsInNjb3BlcyI6WyJhcGlfcmVhZCJdLCJ2ZXJzaW9uIjoxfQ.mniwNJ0_EVyiz75FjAgZWI2SUFKoTV2A9eci3cpcEhQ"}       
        response = requests.get(url, headers=headers)
        dados_filme = response.json()
        url_base_imagem = "https://image.tmdb.org/t/p/w500"
        poster_path = url_base_imagem + dados_filme['poster_path'] if dados_filme.get('poster_path') else None
        produtoras = [{
            'nome': produtora['name'],
            'logo_path': url_base_imagem + produtora['logo_path'] if produtora.get('logo_path') else None,
            'origem': produtora['origin_country']
        } for produtora in dados_filme.get('production_companies', [])]

        return {'poster_path': poster_path, 'produtoras': produtoras}


    def criar_documento_word(self, movie_id, elenco_json, detalhes_filme, atores_info):
        pasta_filme = str(movie_id)
        if not os.path.exists(pasta_filme):
            os.makedirs(pasta_filme)
        
        doc = Document()
        
        doc.add_heading('Detalhes do Filme', 0)
        
        if detalhes_filme['poster_path']:
            response = requests.get(detalhes_filme['poster_path'], stream=True)
            poster_path = os.path.join(pasta_filme, 'poster.jpg')
            if response.status_code == 200:
                with open(poster_path, 'wb') as f:
                    f.write(response.content)
                doc.add_picture(poster_path, width=Inches(2))
                doc.add_paragraph()
        
        doc.add_heading('Top 5 Atores:', level=1)
        for ator in atores_info:
            if ator['url_foto']:
                response = requests.get(ator['url_foto'], stream=True)
                ator_path = os.path.join(pasta_filme, f"{ator['nome'].replace(' ', '_')}.jpg")
                if response.status_code == 200:
                    with open(ator_path, 'wb') as f:
                        f.write(response.content)
                    doc.add_paragraph(ator['nome'])
                    doc.add_picture(ator_path, width=Inches(1))
                    doc.add_paragraph()
        
        doc.add_heading('Produtoras:', level=1)
        for produtora in detalhes_filme['produtoras']:
            if produtora['logo_path']:
                response = requests.get(produtora['logo_path'], stream=True)
                produtora_path = os.path.join(pasta_filme, f"{produtora['nome'].replace(' ', '_')}.png")
                if response.status_code == 200:
                    with open(produtora_path, 'wb') as f:
                        f.write(response.content)
                    doc.add_paragraph(produtora['nome'])
                    doc.add_picture(produtora_path, width=Inches(1))
                    doc.add_paragraph()
        
        doc_path = os.path.join(pasta_filme, 'Detalhes_do_Filme.docx')
        doc.save(doc_path)

    def converter_docx_para_pdf(self, pasta, nome_arquivo):
        # Caminho completo do arquivo .docx
        caminho_docx = os.path.join(pasta, nome_arquivo + '.docx')
        # Caminho completo do arquivo .pdf a ser criado
        caminho_pdf = os.path.join(pasta, nome_arquivo + '.pdf')
        
        # Converter o documento
        convert(caminho_docx, caminho_pdf)
        print(f"Arquivo convertido para PDF: {caminho_pdf}")